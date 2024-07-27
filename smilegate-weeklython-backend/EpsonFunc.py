import dropbox
import fitz  # PyMuPDF
import openai
from langchain_openai import ChatOpenAI
from langchain_community.document_loaders import UnstructuredFileLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import CacheBackedEmbeddings
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma, FAISS
from langchain.storage import LocalFileStore
from langchain.chains import RetrievalQA
from langchain.prompts import ChatPromptTemplate
from langchain.schema.runnable import RunnablePassthrough
from langchain.prompts.chat import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain.agents.format_scratchpad import format_to_openai_function_messages
from langchain.memory import ConversationBufferMemory
from langchain.agents import create_openai_functions_agent, AgentExecutor
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.schema import StrOutputParser
import json
import os
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.pydantic_v1 import BaseModel, Field
from typing import List
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx.shared import Inches
from langchain.tools.retriever import create_retriever_tool
from langchain_community.agent_toolkits.load_tools import load_tools
from langchain import hub
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
import io
from dropbox import Dropbox
from dropbox.files import WriteMode
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from langchain_core.agents import AgentActionMessageLog, AgentFinish
import openai
from dotenv import load_dotenv
load_dotenv()

openai_api_key = os.environ.get('OPENAI_API_KEY')
openai_api_key2 = os.environ.get('OPENAI_API_KEY2')

# Set your Dropbox access token and folder path
ACCESS_TOKEN = "sl.B5w2mg28UCZuyXAINCedxiBbHZ07V9rYIBtux2uAg5t1uX-r18TrVetIbqHKJJSnm9VBSFnaxYmPhdEUD8cq331laFY5zyujreEFyOdiKEEsiCWEbV_LVsiSnfXOg_HsrE77NSheneqif6Knq_DN4R4"



def MakeStudentInfo(Question_pdf, Answer_json):
    #document = fitz.open(Question_pdf)
    document = Question_pdf
    # 텍스트를 저장할 파일 열기
    Question_pdf_file = "Epson-data/Question_pdf.txt"
    with open(Question_pdf_file, 'w', encoding='utf-8') as file:
        # PDF 파일의 각 페이지에서 텍스트 추출
        for page_number in range(len(document)):
            page = document.load_page(page_number)
            text = page.get_text()
            file.write(f"Page {page_number + 1}:\n")
            file.write(text)
            file.write("\n\n")
    
    document.close()
    with open(Question_pdf_file, 'r', encoding='utf-8') as file:
        extracted_text = file.read()

    # JSON 파일 읽기
    #with open(Answer_json, 'r', encoding='utf-8') as file:
    #    json_data = json.load(file)

    # Pydantic 모델 정의
    class StudentResponse(BaseModel):
        Number: int = Field(description="문제 번호")
        Question: str = Field(description="문제")
        Text: str = Field(description="문제에 첨부된 지문")
        Candidate: List[str] = Field(description="5가지 정답 후보")
        StudentAnswer: int = Field(description="학생이 제출한 정답")

    llm = ChatOpenAI(openai_api_key=openai_api_key, 
                    temperature=0.1, 
                    model="gpt-4o")
    
    human_message_prompt = """
    {extracted_text} 이것은 N개의 문제이며, 
    {Answer_sheet_json} 이것은 N개의 문제에 대해 제출한 학생의 답안지이다. 
    {Answer_sheet_json} 위 파일을 보면, 몇개의 문제인지 인식 가능하다. 
    N개의 문제를 인식해서 문제 번호, 문제, 보기, 학생이 제출한 정답을 모두 출력해줘.
    Generate a JSON object with the following fields:
    - Number: 각 문제 번호
    - Question: 각 문제
    - Text: 각 문제에 첨부된 지문
    - Candidate: 각 문제에 대한 보기
    - StudentAnswer: 학생이 제출한 정답

    Here is a Example:
    - Number: 1
    - Question: '<보기>를 바탕으로 접사에 대해 탐구한 내용으로 적절하지 않은 것은?'
    - Text: '<보기>\n접두사와 접미사는 다양한 품사의 어근에 결합하여 파생어를 형성한다. 접두사는 어근에 어휘 적 의미를 더해 주는 경우가 많으며, 파생어의 품사를 결정하는 경우는 거의 없다. 한편 접미사는 파생 과정에서 어근에 문법적 의미를 더해 주는 경우가 있으며, 파생어의 품사가 어근의 품사와 달라지는 경우가 있다.'
    - Candidate: ["① '덧저고리, 덧대다'를 보니 접두사 '덧-'은 명사에도 결합하고 동사에도 결합하는구나.", 
                  "② '군말, 군살'을 보니 접두사 '군-'은 '쓸데없는'이라는 어휘적 의미를 어근에 더해 주는구나.", 
                  "③ '잠꾸러기, 욕심꾸러기'를 보니 어근에 접미사 '-꾸러기'가 결합하면 파생어의 품사가 달라 지는구나.", 
                  "④ '늦추다, 낮추다'를 보니 접미사 '-추-'는 어근 '늦-'과 '낮-'에 사동이라는 문법적 의미를 더해 주는구나.", 
                  "⑤ '풋과일, 헛디디다'를 보니 접두사 '풋-'과 '헛-'이 어근에 결합해도 파생어의 품사는 어근의 품사와 달라지지 않는구나."]
    - StudentAnswer: 3"""

    parser = JsonOutputParser(pydantic_object=StudentResponse)
    #format_instructions = parser.get_format_instructions()
    chat_prompt_template = ChatPromptTemplate.from_messages([human_message_prompt])
    chain = chat_prompt_template | llm | parser
    response = chain.invoke({"extracted_text": extracted_text, 
                             "Answer_sheet_json": Answer_json})
    # JSON 파일로 저장
    with open('./QuestionAnswer.json', 'w', encoding='utf-8') as json_file:
        json.dump(response, json_file, ensure_ascii=False, indent=4)
    
    
    return response

def MakeScoreCommentary(QuestionAnswerJson):

    # JSON 파일 열기
    #with open(QuestionAnswerJson, 'r') as file:
        # JSON 파일 읽기 및 파싱
        #data = json.load(file)
    data = QuestionAnswerJson
        # Pydantic 모델 정의
    class CommentaryResponse(BaseModel):
        Number: int = Field(description="문제 번호")
        StudentAnswer: int = Field(description="학생이 제출한 정답")
        CorrectAnswer: int = Field(description="실제 정답")
        IsCorrect: bool = Field(description="정답 여부")
        CommentarySummarize: List[str] = Field(description="문제 해설 3줄 요약")
    
    cache_dir = LocalFileStore("./.cache/")

    splitter = CharacterTextSplitter.from_tiktoken_encoder(
        separator="\n",
        chunk_size=600,
        chunk_overlap=100,
    )
    loader = UnstructuredFileLoader("data/2024 수능특강 언어와 매체 정답.pdf")

    docs = loader.load_and_split(text_splitter=splitter)

    embeddings = OpenAIEmbeddings()

    cached_embeddings = CacheBackedEmbeddings.from_bytes_store(embeddings, cache_dir)

    vectorstore = FAISS.from_documents(docs, cached_embeddings)

    retriver = vectorstore.as_retriever()
    
    
    llm = ChatOpenAI(openai_api_key=openai_api_key2, 
                 model= "gpt-4o", 
                 temperature=0.1)
    structured_llm = llm.with_structured_output(CommentaryResponse)
    # parser가 get_format_instructions를 통해 스키마 생성, LLM 출력을 Json 형식으로 파싱
    parser = JsonOutputParser(pydantic_object=CommentaryResponse)
    format_instructions = parser.get_format_instructions()
    prompt = PromptTemplate(
    template = """Answer the question based only on the following context:
    {context}

    Question: {question}\n{format_instructions}
    """,
        input_variables=["context", "question"],
        partial_variables={"format_instructions": format_instructions},
    )
    chain = (
        {
            "context": retriver,
            "question": RunnablePassthrough(),
        }
        | prompt
        | structured_llm
        #| parser
    )
    promptforAnswer = """
    위 내용은 문제, 보기, 학생이 제출한 답안이다. 
    정답지에서 ① = 1, ② = 2, ③ = 3, ④ = 4, ⑤ = 5 을 의미한다.
    정답지를 활용해 문제 번호, 학생이 제출한 답안, ,실제 정답, 학생의 정답 여부을 작성해줘.
    문제해설을 작성해주되, 학생의 답안은 정확하다. 와 같은 해설은 보여주지 말아줘. 3줄로 요약해줘.
    아래 예시처럼 작성해줘.
    - Number: 1
    - StudentAnswer: 2
    - CorrectAnswer: 5
    - IsCorrect: False
    - CommentarySummarize: 
    ["토론의 내용은 아동의 개인 정보 노출과 관련된 법적 보호 필요성에 초점을 맞추고 있음.",
    "찬성 측은 개인 정보 보호의 중요성을 강조하며 법적 보호를 주장할 수 있음.",
    "반대 측은 개인 정보의 미숙한 권리 행사나 삭제 요구 권리의 제도화를 주장할 수 있음."]
    """

    Question_list = []
    
    # 각 객체의 'name' 값을 리스트에 추가
    
    for item in data:
        Question = f"문제 번호: {item['Number']} \n문제: {item['Question']} \n지문: {item['Text']} \n보기: {item['Candidate']} \n학생이 제출한 답안: {item['StudentAnswer']}\n" + promptforAnswer
        Question_list.append(Question)

    
    response = chain.batch(Question_list)
    
    #print(Total_response)
    response_list = []
    for resp in response:
        resp_dict = {"Number": resp.Number,
                     "StudentAnswer": resp.StudentAnswer,
                     "CorrectAnswer": resp.CorrectAnswer,
                     "IsCorrect": resp.IsCorrect,
                     "CommentarySummarize": resp.CommentarySummarize}
        response_list.append(resp_dict)
    
    # JSON 파일로 저장
    with open('./AnswerCommentary.json', 'w', encoding='utf-8') as json_file:
        json.dump(response_list, json_file, ensure_ascii=False, indent=4)
    
    
    return response_list


def CreateMemorizationBookAddCommentary(QuestionAnswer_json, AnswerCommentary_json, Output_path):
    # JSON 파일 열기
    with open(QuestionAnswer_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        QuestionAnswer_json_data = json.load(file)

    with open(AnswerCommentary_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        AnswerCommentary_json_data = json.load(file)

    doc = Document()
    dbx = dropbox.Dropbox(ACCESS_TOKEN)
    for i in range(len(QuestionAnswer_json_data)):
        if AnswerCommentary_json_data[i]['IsCorrect'] == False:
            
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(5)

            # 날짜, 월, 일
            now = datetime.now()
            current_date = now.date()
            cell = table.cell(0, 0)
            cell.text = "날짜"
            cell = table.cell(0, 1)
            cell.text = f"{current_date.year}년 {current_date.month}월 {current_date.day}일"
            
            # 과목
            cell = table.cell(1, 0)
            cell.text = "과목"
            cell = table.cell(1, 1)
            cell.text = "언어와 매체"

            cell = table.cell(2, 0)
            cell.text = "문제"
            cell = table.cell(2, 1)
            # 문제 + 지문 + 보기 5개 + 정답 + 해설
            cell.text = QuestionAnswer_json_data[i]['Question'] + "\n\n" + QuestionAnswer_json_data[i]['Text']

            cell = table.cell(3, 0)
            cell.text = "보기"
            cell = table.cell(3, 1)
            cell.text = "\n".join(QuestionAnswer_json_data[i]['Candidate'])
            # 난이도
            cell = table.cell(4, 0)
            cell.text = "정답"
            cell = table.cell(4, 1)
            cell.text = str(AnswerCommentary_json_data[i]['CorrectAnswer'])
            cell = table.cell(5, 0)
            cell.text = "해설"
            cell = table.cell(5, 1)
            cell.text = "\n".join(AnswerCommentary_json_data[i]['CommentarySummarize'])

            doc.add_page_break()
    
    #doc.save(Output_path)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    file_path = Output_path  # Dropbox에 저장될 경로와 파일명
    try:
        dbx.files_upload(doc_stream.getvalue(), file_path, mode=WriteMode('overwrite'))
        print(f"파일이 성공적으로 업로드되었습니다: {file_path}")
    except Exception as e:
        print(f"업로드 중 오류가 발생했습니다: {str(e)}")
    # 단일 파일 변환
    # 사용 예제

def CreateMemorizationBook(QuestionAnswer_json, AnswerCommentary_json, Output_path):
    # JSON 파일 열기
    with open(QuestionAnswer_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        QuestionAnswer_json_data = json.load(file)

    with open(AnswerCommentary_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        AnswerCommentary_json_data = json.load(file)

    dbx = dropbox.Dropbox(ACCESS_TOKEN)
    doc = Document()
    
    for i in range(len(QuestionAnswer_json_data)):
        if AnswerCommentary_json_data[i]['IsCorrect'] == False:
            
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(5)

            # 날짜, 월, 일
            now = datetime.now()
            current_date = now.date()
            cell = table.cell(0, 0)
            cell.text = "날짜"
            cell = table.cell(0, 1)
            cell.text = f"{current_date.year}년 {current_date.month}월 {current_date.day}일"
            
            # 과목
            cell = table.cell(1, 0)
            cell.text = "과목"
            cell = table.cell(1, 1)
            cell.text = "언어와 매체"

            cell = table.cell(2, 0)
            cell.text = "문제"
            cell = table.cell(2, 1)
            # 문제 + 지문 + 보기 5개 + 정답 + 해설
            cell.text = QuestionAnswer_json_data[i]['Question'] + "\n\n" + QuestionAnswer_json_data[i]['Text']

            cell = table.cell(3, 0)
            cell.text = "보기"
            cell = table.cell(3, 1)
            cell.text = "\n".join(QuestionAnswer_json_data[i]['Candidate'])
            # 난이도
            cell = table.cell(4, 0)
            cell.text = "정답"
            cell = table.cell(4, 1)
            cell.text = ""
            cell = table.cell(5, 0)
            cell.text = "해설"
            cell = table.cell(5, 1)
            cell.text = ""

            doc.add_page_break()
    
    #doc.save(Output_path)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    file_path = Output_path  # Dropbox에 저장될 경로와 파일명
    try:
        dbx.files_upload(doc_stream.getvalue(), file_path, mode=WriteMode('overwrite'))
        print(f"파일이 성공적으로 업로드되었습니다: {file_path}")
    except Exception as e:
        print(f"업로드 중 오류가 발생했습니다: {str(e)}")
    # 단일 파일 변환
    # 사용 예제

def load_pdfs_from_folder(folder_path):
    total_split_docs = []
    # 텍스트 분할기를 사용하여 문서를 분할합니다.
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

    for filename in os.listdir(folder_path):
        if filename.endswith('.pdf'):
            file_path = os.path.join(folder_path, filename)
            loader = PyPDFLoader(file_path)
            split_docs = loader.load_and_split(text_splitter=text_splitter)
            total_split_docs.extend(split_docs)
    return total_split_docs

def parse(output):
    # If no function was invoked, return to user
    if "function_call" not in output.additional_kwargs:
        return AgentFinish(return_values={"output": output.content}, log=output.content)

    # Parse out the function call
    function_call = output.additional_kwargs["function_call"]
    name = function_call["name"]
    inputs = json.loads(function_call["arguments"])

    # If the Response function was invoked, return to the user with the function inputs
    if name == "Response":
        return AgentFinish(return_values=inputs, log=str(function_call))
    # Otherwise, return an agent action
    else:
        return AgentActionMessageLog(
            tool=name, tool_input=inputs, log="", message_log=[output]
        )

def CreateCorrectAnswerNote(QuestionAnswer_json, AnswerCommentary_json, Output_path):
    with open(QuestionAnswer_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        QuestionAnswer_json_data = json.load(file)

    with open(AnswerCommentary_json, 'r') as file:
        # JSON 파일 읽기 및 파싱
        AnswerCommentary_json_data = json.load(file)

    folder_path = 'data/Insendium'
    pages = load_pdfs_from_folder(folder_path)
    dbx = dropbox.Dropbox(ACCESS_TOKEN)
    # VectorStore를 생성합니다.
    vector = FAISS.from_documents(pages, OpenAIEmbeddings())

    # Retriever를 생성합니다.
    retriever = vector.as_retriever()
    
    # hub에서 prompt를 가져옵니다 - 이 부분을 수정할 수 있습니다!
    #prompt = hub.pull("hwchase17/openai-functions-agent")
    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", "You are a helpful assistant"),
            ("user", "{input}"),
            MessagesPlaceholder(variable_name="agent_scratchpad"),
        ]
    )
    # Pydantic 모델 정의
    class AgentResponse(BaseModel):
        Question: str = Field(description="문제")
        Text: str = Field(description="문제에 첨부된 지문")
        Candidate: List[str] = Field(description="5가지 정답 후보")
        CorrectAnswer: int = Field(description="문제 정답")
        CommentarySummarize: List[str] = Field(description="문제 해설 3줄 요약")

    llm = ChatOpenAI(model="gpt-4-turbo", temperature=0)
    # parser가 get_format_instructions를 통해 스키마 생성, LLM 출력을 Json 형식으로 파싱
    parser = JsonOutputParser(pydantic_object=AgentResponse)
    format_instructions = parser.get_format_instructions()
    doc = Document()
    
    for i in range(len(QuestionAnswer_json_data)):
        if AnswerCommentary_json_data[i]['IsCorrect'] == False:
            Question_Text = QuestionAnswer_json_data[i]['Question'] + QuestionAnswer_json_data[i]['Text']
            # langchain 패키지의 tools 모듈에서 retriever 도구를 생성
            retriever_tool = create_retriever_tool(
                retriever=retriever,
                name="pdf_search",
                # 도구에 대한 설명을 자세히 기입해야 합니다!!!
                description= Question_Text + "와 유사한 문서를 PDF 문서에서 검색합니다. 이외 문서는 검색하지 않는다."
            )
            search = load_tools(["ddg-search"])[0]
            # tools 리스트에 search와 retriever_tool을 추가합니다.
            tools = [search, retriever_tool]
            #structured_llm = llm.with_structured_output(AgentResponse)
            llm_with_tools = llm.bind_functions([search, retriever_tool])
            
            # llm, tools, prompt를 인자로 사용합니다.
            # agent = create_openai_functions_agent(llm, tools, prompt)
            agent = (
                {
                    "input": lambda x: x["input"],
                    # Format agent scratchpad from intermediate steps
                    "agent_scratchpad": lambda x: format_to_openai_function_messages(
                        x["intermediate_steps"]
                    ),
                }
                | prompt
                | llm_with_tools
                | parse
            )
            # AgentExecutor 클래스를 사용하여 agent와 tools를 설정하고, 상세한 로그를 출력하도록 verbose를 True로 설정합니다.
            agent_executor = AgentExecutor(tools = tools, agent=agent, verbose=True)

            
            response = agent_executor.invoke(
                {"input": Question_Text + """과 관련성이 높은 문제를 PDF 문서에서 찾고, 관련성이 높은 문제를 참고해서 유사한 문제를 만들어줘. 또한, 해설과 정답까지 알려줘. 결과를 json object로 출력해줘.
                            Generate the answer in Json format.
                            Here is the Json Output example.
                            {"Question": "밑줄 친 부분이 <보기>의 ㉠~㉤에 해당하는 예로 적절한 것은?" ,
                            "Text": "<보기>\n안은문장은 한 절이 그 속에 다른 절을 문장 성분의 하나로 안고 있는 문장이다. 이때 안겨 있는 절을 안긴절이라 하며, 안긴절의 종류에는 ㉠명사절, ㉡관형사절, ㉢부사절, ㉣서술절, ㉤인용절이 있다. 명사절, 관형사절, 부사절은 주로 전성 어미를 통해 실현된다. 서술절은 전성 어미 없이 실현되며, 인용절은 조사가 붙어 실현된다.", 
                            "Candidate": ["1.㉠: 그가 학교에 간다는 사실을 알고 있었다.",
                                        "2.㉡: 내가 어제 본 영화는 매우 재미있었다.",
                                        "3.㉢: 비가 오기 전에 집에 도착해야 한다.",
                                        "4.㉣: 그녀가 도시에 도착했음을 확인했다.",
                                        "5.㉤: 그는 내일 회사를 그만둔다고 말했다."],
                            "CorrectAnswer": 5, 
                            "CommentarySummarize": 
                            ["토론의 내용은 아동의 개인 정보 노출과 관련된 법적 보호 필요성에 초점을 맞추고 있음.",
                            "찬성 측은 개인 정보 보호의 중요성을 강조하며 법적 보호를 주장할 수 있음.",
                            "반대 측은 개인 정보의 미숙한 권리 행사나 삭제 요구 권리의 제도화를 주장할 수 있음."]}\{format_instructions}
                            """,
                            },
                partial_variables={"format_instructions": format_instructions},
                return_only_outputs=True,
            )
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(5)

            # 날짜, 월, 일
            now = datetime.now()
            current_date = now.date()
            cell = table.cell(0, 0)
            cell.text = "날짜"
            cell = table.cell(0, 1)
            cell.text = f"{current_date.year}년 {current_date.month}월 {current_date.day}일"
            
            # 과목
            cell = table.cell(1, 0)
            cell.text = "과목"
            cell = table.cell(1, 1)
            cell.text = "언어와 매체"

            cell = table.cell(2, 0)
            cell.text = "문제"
            cell = table.cell(2, 1)
            # 문제 + 지문 + 보기 5개 + 정답 + 해설
            resp = response['output']
            data = json.loads(resp)
            cell.text = data['Question'] + "\n\n" + data['Text']
            cell = table.cell(3, 0)
            cell.text = "보기"
            cell = table.cell(3, 1)
            cell.text = "\n".join(data['Candidate'])
            # 난이도
            cell = table.cell(4, 0)
            cell.text = "정답"
            cell = table.cell(4, 1)
            cell.text = str(data['CorrectAnswer'])
            cell = table.cell(5, 0)
            cell.text = "해설"
            cell = table.cell(5, 1)
            cell.text = "\n".join(data['CommentarySummarize'])

            doc.add_page_break()
    
    #doc.save(Output_path)
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    file_path = Output_path  # Dropbox에 저장될 경로와 파일명
    try:
        dbx.files_upload(doc_stream.getvalue(), file_path, mode=WriteMode('overwrite'))
        print(f"파일이 성공적으로 업로드되었습니다: {file_path}")
    except Exception as e:
        print(f"업로드 중 오류가 발생했습니다: {str(e)}")

def MakeStudentInfoScoreCommentary(Question_pdf, Answer_json):
    response1 = MakeStudentInfo(Question_pdf = Question_pdf,
                Answer_json = Answer_json)
    
    response2 = MakeScoreCommentary(QuestionAnswerJson = response1)
    return response2
